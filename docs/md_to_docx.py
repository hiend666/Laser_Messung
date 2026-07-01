"""
md_to_docx.py – Wandelt docs/Anleitung.md in docs/Anleitung.docx um.

Erzeugt eine echte Word-Datei, die in MS Word, LibreOffice Writer oder
Google Docs nachbearbeitet werden kann. Nutzt:
  - Heading-Styles (Heading 1/2/3) -> navigation / Inhaltsverzeichnis-fähig
  - Aufzählungen (List Bullet)
  - Word-Tabellen mit dem Tabellen-Style "Light Grid Accent 1"
  - zentrierte Bilder mit kursivem Caption (fett vs. kursiv wird unterstützt)

Abhängigkeit: python-docx (gerade in das venv installiert).
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Inline-Markdown: liefert (text, runs)-Beschreibung und wendet **bold** / *italic* / `code` an
# ---------------------------------------------------------------------------
def add_inline(paragraph, text: str):
    """Fügt dem Paragraph mehrere Runs mit Inline-Formatierung hinzu."""
    # Token-basierter Scanner: erkennt `code`, **bold**, *italic*
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------------------------------------------------------------------------
# Tabelle aus Markdown-Zeilen bauen
# ---------------------------------------------------------------------------
def add_table(doc, tbl_lines: list):
    rows = []
    for ln in tbl_lines:
        if re.match(r"^\|[\s:|-]+\|\s*$", ln):
            continue  # Trennzeile
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < n_cols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass  # Style nicht vorhanden -> Standard
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Zelle leeren (Default-Paragraph vorhanden)
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            if r_idx == 0:
                # Header fett
                run = p.add_run(_strip_inline(cell_text))
                run.bold = True
            else:
                add_inline(p, cell_text)


def _strip_inline(text: str) -> str:
    """Entfernt Markdown-Zeichen für Tabellen-Header (einfach fett)."""
    s = re.sub(r"`([^`]+)`", r"\1", text)
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"\*([^*]+)\*", r"\1", s)
    return s


# ---------------------------------------------------------------------------
# Bild mit proportionaler Skalierung einfügen
# ---------------------------------------------------------------------------
def add_image(doc, img_path: Path, max_w_cm: float = 16.0):
    """Fügt ein zentriertes, proportional skaliertes Bild ein."""
    from PIL import Image as PILImage
    try:
        with PILImage.open(img_path) as im:
            iw, ih = im.size
    except Exception:
        iw, ih = 1000, 600
    # Bildbreite; Höhe proportional, max 18 cm
    w = Cm(max_w_cm)
    ratio = ih / iw
    h = Cm(max_w_cm * ratio)
    if h > Cm(18.0):
        scale = 18.0 / (max_w_cm * ratio)
        w = Cm(max_w_cm * scale)
        h = Cm(18.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(img_path), width=w, height=h)
    except Exception as e:
        p.add_run(f"[Bild-Fehler: {e}]")
    return p


# ---------------------------------------------------------------------------
# Markdown-Parser -> Word
# ---------------------------------------------------------------------------
def is_block_start(line: str) -> bool:
    s = line.rstrip()
    return (
        s.startswith("|")
        or re.match(r"^#{1,3}\s", s)
        or re.match(r"^!\[", s)
        or re.match(r"^\s*[-*]\s+", s)
        or re.match(r"^-{3,}$", s)
        or s.startswith(">")
    )


def parse_to_docx(md_text: str, base_dir: Path, out_path: Path):
    doc = Document()

    # Basis-Stile: Schriftfarbe für Überschriften etwas absetzen
    styles = doc.styles
    try:
        styles["Heading 1"].font.color.rgb = RGBColor(0x1a, 0x3a, 0x6c)
        styles["Heading 2"].font.color.rgb = RGBColor(0x1a, 0x3a, 0x6c)
        styles["Heading 3"].font.color.rgb = RGBColor(0x2a, 0x4a, 0x7c)
        styles["Normal"].font.size = Pt(10.5)
    except Exception:
        pass

    lines = md_text.splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        raw = line.rstrip()

        if not raw.strip():
            i += 1
            continue

        # Trennlinie
        if re.match(r"^-{3,}$", raw):
            p = doc.add_paragraph()
            # Horizontale Rule via unterer Rahmen
            pPr = p.paragraph_format
            try:
                from docx.oxml import OxmlElement
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "999999")
                pBdr.append(bottom)
                p.paragraph_format.element.get_or_add_pPr().append(pBdr)
            except Exception:
                pass
            i += 1
            continue

        # Überschriften
        m = re.match(r"^(#{1,3})\s+(.*)$", raw)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading(level=level)
            add_inline(h, text)
            i += 1
            continue

        # Bild
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", raw)
        if m:
            alt, rel = m.group(1), m.group(2)
            img_path = (base_dir / rel).resolve()
            if not img_path.exists():
                doc.add_paragraph(f"[Bild fehlt: {rel}]")
                i += 1
                continue
            add_image(doc, img_path)
            if alt:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(alt)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Tabelle
        if raw.startswith("|"):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            add_table(doc, tbl_lines)
            continue

        # Zitat
        if raw.startswith(">"):
            quote_lines = []
            while i < n and lines[i].lstrip().startswith(">"):
                quote_lines.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            quote_text = " ".join(l.strip() for l in quote_lines if l.strip())
            if quote_text:
                qp = doc.add_paragraph()
                qp.paragraph_format.left_indent = Cm(0.6)
                run = qp.add_run(_strip_inline(quote_text))
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Aufzählung
        if re.match(r"^\s*[-*]\s+", raw):
            items = []
            while i < n and re.match(r"^\s*[-*]\s+", lines[i]):
                item = re.sub(r"^\s*[-*]\s+", "", lines[i]).strip()
                items.append(item)
                i += 1
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, item)
            continue

        # Absatz
        block = []
        while i < n and lines[i].strip() and not is_block_start(lines[i]):
            block.append(lines[i])
            i += 1
        text = " ".join(l.strip() for l in block if l.strip())
        if text:
            p = doc.add_paragraph()
            add_inline(p, text)

    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    base_dir = Path(__file__).resolve().parent.parent / "docs"
    md_path = base_dir / "Anleitung.md"
    out_path = base_dir / "Anleitung.docx"
    if not md_path.exists():
        print(f"FEHLER: {md_path} nicht gefunden", file=sys.stderr)
        sys.exit(1)
    md_text = md_path.read_text(encoding="utf-8")
    parse_to_docx(md_text, base_dir, out_path)
    print(f"DOCX erstellt: {out_path}")


if __name__ == "__main__":
    main()
